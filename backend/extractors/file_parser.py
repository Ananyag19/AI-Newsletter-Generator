"""
Extracts text from uploaded files: PDF, DOCX, and plain text.
"""
import io
import logging

import fitz  # PyMuPDF
from docx import Document

from models.schemas import ExtractedContent

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


def _parse_pdf(file_bytes: bytes, filename: str) -> ExtractedContent:
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages_text = []
        for page in doc:
            pages_text.append(page.get_text())
        text = "\n".join(pages_text).strip()
        title = doc.metadata.get("title") if doc.metadata else None
        doc.close()
        return ExtractedContent(source=filename, title=title or None, text=text)
    except Exception as e:  # noqa: BLE001
        logger.exception("Failed to parse PDF %s", filename)
        return ExtractedContent(source=filename, text="", error=f"Failed to parse PDF: {e}")


def _parse_docx(file_bytes: bytes, filename: str) -> ExtractedContent:
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also pull table cell text since notes/docs often use tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())

        text = "\n".join(paragraphs).strip()
        title = paragraphs[0] if paragraphs else None
        return ExtractedContent(source=filename, title=title, text=text)
    except Exception as e:  # noqa: BLE001
        logger.exception("Failed to parse DOCX %s", filename)
        return ExtractedContent(source=filename, text="", error=f"Failed to parse DOCX: {e}")


def _parse_txt(file_bytes: bytes, filename: str) -> ExtractedContent:
    try:
        text = file_bytes.decode("utf-8", errors="replace").strip()
        return ExtractedContent(source=filename, text=text)
    except Exception as e:  # noqa: BLE001
        return ExtractedContent(source=filename, text="", error=f"Failed to read text file: {e}")


def parse_file(file_bytes: bytes, filename: str) -> ExtractedContent:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return _parse_pdf(file_bytes, filename)
    if lower.endswith(".docx"):
        return _parse_docx(file_bytes, filename)
    if lower.endswith(".txt") or lower.endswith(".md"):
        return _parse_txt(file_bytes, filename)
    return ExtractedContent(
        source=filename,
        text="",
        error=f"Unsupported file type. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}",
    )
